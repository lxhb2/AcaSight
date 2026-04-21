#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word报告生成功能
自动生成XRD分析报告
"""

from pathlib import Path
from typing import Dict, Any, Optional
import json
from datetime import datetime

class WordReportGenerator:
    """Word报告生成器"""
    
    def __init__(self):
        self.template_dir = Path(__file__).parent / "templates"
        self.template_dir.mkdir(exist_ok=True)
    
    def generate_report(self, analysis_data: Dict[str, Any], 
                       output_path: str,
                       template: str = "default") -> Dict[str, Any]:
        """
        生成Word报告
        
        参数:
            analysis_data: XRD分析数据
            output_path: 输出文件路径
            template: 模板名称
        
        返回:
            生成结果信息
        """
        try:
            output_path = Path(output_path)
            
            # 根据模板选择生成方法
            if template == "default":
                return self._generate_default_report(analysis_data, output_path)
            elif template == "detailed":
                return self._generate_detailed_report(analysis_data, output_path)
            elif template == "simple":
                return self._generate_simple_report(analysis_data, output_path)
            else:
                # 尝试加载自定义模板
                template_file = self.template_dir / f"{template}.json"
                if template_file.exists():
                    return self._generate_from_template(analysis_data, output_path, template_file)
                else:
                    raise ValueError(f"模板不存在: {template}")
                    
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "template": template,
                "output_path": str(output_path)
            }
    
    def _generate_default_report(self, data: Dict[str, Any], output_path: Path) -> Dict[str, Any]:
        """生成默认报告"""
        try:
            # 创建报告内容
            report_content = self._create_report_content(data, "default")
            
            # 保存为文本文件（实际应用中可以使用python-docx库生成真正的Word文档）
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            # 如果安装了python-docx，可以生成真正的Word文档
            try:
                import docx
                return self._generate_docx_report(data, output_path.with_suffix('.docx'))
            except ImportError:
                # 回退到文本格式
                pass
            
            return {
                "success": True,
                "format": "txt",
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "template": "default"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "template": "default",
                "output_path": str(output_path)
            }
    
    def _create_report_content(self, data: Dict[str, Any], template: str) -> str:
        """创建报告内容"""
        metadata = data.get('metadata', {})
        file_info = data.get('file_info', {})
        analysis_results = data.get('analysis_results', {})
        ai_analysis = data.get('ai_analysis', {})
        summary = data.get('summary', {})
        
        # 报告标题
        report = f"""XRD分析报告
{'=' * 60}

报告信息:
{'=' * 60}
生成时间: {metadata.get('analysis_time', datetime.now().isoformat())}
分析文件: {file_info.get('filename', 'N/A')}
文件大小: {file_info.get('size', 0):,} 字节
数据点数: {file_info.get('data_points', 0)}

分析参数:
{'=' * 60}
"""
        
        # 分析参数
        params = metadata.get('parameters', {})
        for key, value in params.items():
            report += f"{key}: {value}\n"
        
        # 分析结果摘要
        report += f"""
分析结果摘要:
{'=' * 60}
检测峰数: {summary.get('peak_count', 0)}
匹配物相数: {summary.get('phase_count', 0)}
主要物相: {summary.get('top_phase', '无匹配')}
匹配置信度: {summary.get('confidence', 0):.3f}
处理时间: {data.get('processing_time', 0):.2f} 秒

"""
        
        # 峰检测结果
        if 'peaks' in analysis_results:
            peaks = analysis_results['peaks']
            report += f"""峰检测结果:
{'=' * 60}
检测到 {len(peaks.get('angles', []))} 个峰

最强峰:
"""
            if peaks.get('angles'):
                strongest_idx = peaks['intensities'].index(max(peaks['intensities'])) if peaks.get('intensities') else 0
                report += f"  角度: {peaks['angles'][strongest_idx]:.2f}°\n"
                report += f"  强度: {peaks['intensities'][strongest_idx]:.2f}\n"
                report += f"  d值: {peaks['d_values'][strongest_idx]:.4f} Å\n"
                report += f"  FWHM: {peaks['fwhms'][strongest_idx]:.4f}°\n"
            
            report += "\n"
        
        # 物相匹配结果
        if 'matched_phases' in analysis_results:
            phases = analysis_results['matched_phases']
            report += f"""物相匹配结果:
{'=' * 60}
匹配到 {len(phases)} 个物相

前5个匹配:
"""
            for i, phase in enumerate(phases[:5], 1):
                report += f"{i}. {phase.get('name', 'N/A')}\n"
                report += f"   化学式: {phase.get('formula', 'N/A')}\n"
                report += f"   卡片号: {phase.get('card_num', 'N/A')}\n"
                report += f"   匹配d值: {phase.get('matched_d', 0):.4f} Å\n"
                report += f"   匹配分数: {phase.get('match_score', 0):.3f}\n"
                report += f"   类型: {phase.get('card_type', 'N/A')}\n\n"
        
        # 统计信息
        if 'statistics' in analysis_results:
            stats = analysis_results['statistics']
            report += f"""统计信息:
{'=' * 60}
角度范围: {stats.get('angle_range', {}).get('min', 0):.1f}-{stats.get('angle_range', {}).get('max', 0):.1f}°
最大强度: {stats.get('intensity_stats', {}).get('max', 0):.2f}
平均强度: {stats.get('intensity_stats', {}).get('mean', 0):.2f}
强度标准差: {stats.get('intensity_stats', {}).get('std', 0):.2f}

"""
        
        # AI分析结果
        if ai_analysis and not ai_analysis.get('fallback', False):
            report += f"""AI分析:
{'=' * 60}
"""
            if 'analysis' in ai_analysis:
                report += f"{ai_analysis['analysis']}\n\n"
            if 'recommendations' in ai_analysis:
                report += "建议:\n"
                for rec in ai_analysis['recommendations']:
                    report += f"• {rec}\n"
                report += "\n"
        
        # 建议
        if 'recommendations' in data:
            report += f"""分析建议:
{'=' * 60}
"""
            for rec in data['recommendations']:
                report += f"• {rec}\n"
            report += "\n"
        
        # 报告结尾
        report += f"""报告结束
{'=' * 60}
生成系统: Sci-XRD v2.0.0
技术支持: QClaw AI Assistant
生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        return report
    
    def _generate_detailed_report(self, data: Dict[str, Any], output_path: Path) -> Dict[str, Any]:
        """生成详细报告"""
        try:
            # 创建更详细的内容
            report_content = self._create_detailed_content(data)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            return {
                "success": True,
                "format": "txt",
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "template": "detailed"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "template": "detailed",
                "output_path": str(output_path)
            }
    
    def _create_detailed_content(self, data: Dict[str, Any]) -> str:
        """创建详细报告内容"""
        # 扩展默认报告，添加更多技术细节
        base_content = self._create_report_content(data, "default")
        
        analysis_results = data.get('analysis_results', {})
        
        detailed_sections = []
        
        # 添加峰详细数据
        if 'peaks' in analysis_results:
            peaks = analysis_results['peaks']
            if peaks.get('angles'):
                detailed_sections.append("峰详细数据:\n" + "=" * 60)
                detailed_sections.append("序号\t角度(°)\t强度\tFWHM(°)\td值(Å)\t突出度")
                
                for i, (angle, intensity, fwhm, d_value) in enumerate(zip(
                    peaks['angles'],
                    peaks['intensities'],
                    peaks.get('fwhms', [0] * len(peaks['angles'])),
                    peaks.get('d_values', [0] * len(peaks['angles']))
                ), 1):
                    prominence = peaks.get('properties', {}).get('prominences', [0] * len(peaks['angles']))[i-1] if i-1 < len(peaks.get('properties', {}).get('prominences', [])) else 0
                    detailed_sections.append(f"{i}\t{angle:.4f}\t{intensity:.2f}\t{fwhm:.4f}\t{d_value:.4f}\t{prominence:.3f}")
        
        # 添加物相详细数据
        if 'matched_phases' in analysis_results:
            phases = analysis_results['matched_phases']
            if phases:
                detailed_sections.append("\n物相详细数据:\n" + "=" * 60)
                detailed_sections.append("序号\t物相名称\t化学式\t卡片号\t匹配d值\t分数\t类型")
                
                for i, phase in enumerate(phases, 1):
                    detailed_sections.append(
                        f"{i}\t{phase.get('name', '')}\t{phase.get('formula', '')}\t"
                        f"{phase.get('card_num', '')}\t{phase.get('matched_d', 0):.4f}\t"
                        f"{phase.get('match_score', 0):.3f}\t{phase.get('card_type', '')}"
                    )
        
        # 组合内容
        detailed_content = base_content + "\n\n" + "\n".join(detailed_sections)
        
        return detailed_content
    
    def _generate_simple_report(self, data: Dict[str, Any], output_path: Path) -> Dict[str, Any]:
        """生成简单报告"""
        try:
            # 创建简洁内容
            metadata = data.get('metadata', {})
            summary = data.get('summary', {})
            
            simple_content = f"""XRD分析简报
{'=' * 40}

文件: {data.get('file_info', {}).get('filename', 'N/A')}
时间: {metadata.get('analysis_time', 'N/A')}

结果:
• 检测峰数: {summary.get('peak_count', 0)}
• 匹配物相: {summary.get('phase_count', 0)}
• 主要物相: {summary.get('top_phase', '无匹配')}
• 置信度: {summary.get('confidence', 0):.1%}

处理时间: {data.get('processing_time', 0):.1f}秒
生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(simple_content)
            
            return {
                "success": True,
                "format": "txt",
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "template": "simple"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "template": "simple",
                "output_path": str(output_path)
            }
    
    def _generate_from_template(self, data: Dict[str, Any], output_path: Path, 
                               template_file: Path) -> Dict[str, Any]:
        """从模板生成报告"""
        try:
            with open(template_file, 'r', encoding='utf-8') as f:
                template_config = json.load(f)
            
            # 应用模板
            report_content = self._apply_template(data, template_config)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            return {
                "success": True,
                "format": "txt",
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "template": template_file.stem
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "template": str(template_file),
                "output_path": str(output_path)
            }
    
    def _apply_template(self, data: Dict[str, Any], template_config: Dict[str, Any]) -> str:
        """应用模板配置"""
        # 这里可以实现更复杂的模板系统
        # 目前使用简单的内容替换
        
        template = template_config.get('template', '')
        
        # 替换变量
        replacements = {
            '{filename}': data.get('file_info', {}).get('filename', ''),
            '{analysis_time}': data.get('metadata', {}).get('analysis_time', ''),
            '{peak_count}': str(data.get('summary', {}).get('peak_count', 0)),
            '{phase_count}': str(data.get('summary', {}).get('phase_count', 0)),
            '{top_phase}': data.get('summary', {}).get('top_phase', ''),
            '{confidence}': f"{data.get('summary', {}).get('confidence', 0):.3f}",
            '{processing_time}': f"{data.get('processing_time', 0):.2f}"
        }
        
        for key, value in replacements.items():
            template = template.replace(key, value)
        
        return template
    
    def _generate_docx_report(self, data: Dict[str, Any], output_path: Path) -> Dict[str, Any]:
        """生成真正的Word文档（需要python-docx库）"""
        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = docx.Document()
            
            # 标题
            title = doc.add_heading('XRD分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 报告信息
            doc.add_heading('报告信息', level=1)
            
            metadata = data.get('metadata', {})
            file_info = data.get('file_info', {})
            
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid'
            
            info_table.cell(0, 0).text = '生成时间'
            info_table.cell(0, 1).text = metadata.get('analysis_time', 'N/A')
            info_table.cell(1, 0).text = '分析文件'
            info_table.cell(1, 1).text = file_info.get('filename', 'N/A')
            info_table.cell(2, 0).text = '文件大小'
            info_table.cell(2, 1).text = f"{file_info.get('size', 0):,} 字节"
            info_table.cell(3, 0).text = '数据点数'
            info_table.cell(3, 1).text = str(file_info.get('data_points', 0))
            
            # 分析结果摘要
            doc.add_heading('分析结果摘要', level=1)
            
            summary = data.get('summary', {})
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid'
            
            summary_table.cell(0, 0).text = '检测峰数'
            summary_table.cell(0, 1).text = str(summary.get('peak_count', 0))
            summary_table.cell(1, 0).text = '匹配物相数'
            summary_table.cell(1, 1).text = str(summary.get('phase_count', 0))
            summary_table.cell(2, 0).text = '主要物相'
            summary_table.cell(2, 1).text = summary.get('top_phase', '无匹配')
            summary_table.cell(3, 0).text = '匹配置信度'
            summary_table.cell(3, 1).text = f"{summary.get('confidence', 0):.3f}"
            
            # 物相匹配结果
            if 'matched_phases' in data.get('analysis_results', {}):
                doc.add_heading('物相匹配结果', level=1)
                
                phases = data['analysis_results']['matched_phases']
                if phases:
                    phase_table = doc.add_table(rows=min(len(phases), 10) +