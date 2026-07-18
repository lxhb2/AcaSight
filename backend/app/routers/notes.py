"""
笔记路由
Markdown 编辑 / 保存 / 导出为 Word
"""

import os
import json
import re
import uuid
from datetime import datetime
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional
import structlog

logger = structlog.get_logger()
router = APIRouter()

NOTES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data", "notes")
os.makedirs(NOTES_DIR, exist_ok=True)

_SAFE_NOTE_ID_PATTERN = re.compile(r"^[a-zA-Z0-9_-]+$")


def _validate_note_id(note_id: str):
    """验证 note_id 只包含安全字符，防止路径遍历"""
    if not _SAFE_NOTE_ID_PATTERN.match(note_id):
        raise HTTPException(400, "note_id 只允许包含字母、数字、连字符和下划线")


class SaveNoteRequest(BaseModel):
    note_id: Optional[str] = None
    content: str
    title: str


class ExportWordRequest(BaseModel):
    markdown: str
    title: str = "note"


@router.post("/save")
async def save_note(req: SaveNoteRequest):
    """保存或更新笔记"""
    note_id = req.note_id or str(uuid.uuid4())
    _validate_note_id(note_id)
    note_file = os.path.join(NOTES_DIR, f"{note_id}.json")
    data = {
        "id": note_id,
        "title": req.title,
        "content": req.content,
        "updated_at": datetime.now().isoformat(),
    }
    with open(note_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)
    return {"id": note_id, "message": "已保存"}


@router.get("/list")
async def list_notes():
    """列举所有笔记"""
    notes = []
    for fname in os.listdir(NOTES_DIR):
        if fname.endswith(".json"):
            try:
                with open(os.path.join(NOTES_DIR, fname), "r", encoding="utf-8") as f:
                    note = json.load(f)
                    notes.append({"id": note.get("id"), "title": note.get("title"), "updated_at": note.get("updated_at")})
            except Exception:
                pass
    notes.sort(key=lambda n: n.get("updated_at", ""), reverse=True)
    return {"notes": notes}


@router.get("/{note_id}")
async def get_note(note_id: str):
    """获取单个笔记"""
    _validate_note_id(note_id)
    note_file = os.path.join(NOTES_DIR, f"{note_id}.json")
    if not os.path.exists(note_file):
        raise HTTPException(404, "笔记不存在")
    with open(note_file, "r", encoding="utf-8") as f:
        return json.load(f)


@router.delete("/{note_id}")
async def delete_note(note_id: str):
    """删除笔记"""
    _validate_note_id(note_id)
    note_file = os.path.join(NOTES_DIR, f"{note_id}.json")
    if os.path.exists(note_file):
        os.remove(note_file)
        return {"message": "已删除"}
    raise HTTPException(404, "笔记不存在")


@router.post("/export-word")
async def export_word(req: ExportWordRequest):
    """Markdown 转 Word (docx)"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx 未安装，请执行: pip install python-docx")

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = req.markdown.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### ') or \
           line.startswith('#### ') or line.startswith('##### ') or line.startswith('###### '):
            level = min(line.count('#', 0, line.index(' ')), 9)
            heading = doc.add_heading(line[level + 1:].strip(), level=level)

        # 无序列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')

        # 有序列表
        elif len(line.strip()) > 2 and line.strip()[0].isdigit() and '. ' in line.strip():
            dot_idx = line.strip().index('. ')
            text = line.strip()[dot_idx + 2:]
            p = doc.add_paragraph(text, style='List Number')

        # 引用
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # 代码块 (简易处理: ``` ... ```)
        elif line.startswith('\'\'\''):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('\'\'\''):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

        # 空行
        elif not line.strip():
            if i > 0 and i < len(lines) - 1 and lines[i - 1].strip() and lines[i + 1].strip():
                doc.add_paragraph()

        # 普通段落
        elif line.strip():
            # 处理行内格式
            p = doc.add_paragraph()
            add_markdown_inline(p, line)

        i += 1

    # 写入内存
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{req.title}.docx"'}
    )


def add_markdown_inline(paragraph, text: str):
    """简易行内 Markdown 渲染: **bold**, *italic*, `code`, [link](url)"""
    import re
    
    # 将 Markdown 转换为 docx 的 run
    patterns = [
        (r'\*\*(.+?)\*\*', lambda m: ('bold', m.group(1))),
        (r'\*(.+?)\*', lambda m: ('italic', m.group(1))),
        (r'`(.+?)`', lambda m: ('code', m.group(1))),
        (r'\[(.+?)\]\((.+?)\)', lambda m: ('link', m.group(1))),
    ]

    # 简化的分词渲染
    i = 0
    remaining = text
    while remaining:
        earliest = (len(remaining), None, None)
        for pat, handler in patterns:
            m = re.search(pat, remaining)
            if m:
                idx = m.start()
                if idx < earliest[0]:
                    earliest = (idx, m, handler)

        if earliest[1] is None:
            # 无更多格式，输出剩余文本
            paragraph.add_run(remaining)
            break

        idx, m, handler = earliest
        if idx > 0:
            paragraph.add_run(remaining[:idx])

        t = handler(m)
        if t[0] == 'bold':
            run = paragraph.add_run(t[1])
            run.bold = True
        elif t[0] == 'italic':
            run = paragraph.add_run(t[1])
            run.italic = True
        elif t[0] == 'code':
            run = paragraph.add_run(t[1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif t[0] == 'link':
            run = paragraph.add_run(t[1])
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True

        remaining = remaining[m.end():]