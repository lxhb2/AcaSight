"""
TemplateService — Word 模板生成服务

支持 GB/T 7714、APA、IEEE 等学术引用格式的 reference-doc.docx 生成。
使用 python-docx 库操作 Word 样式。
"""

import json
import structlog
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = structlog.get_logger()

TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


BUILTIN_TEMPLATES = {
    "gbt7714": {
        "id": "gbt7714",
        "name": "GB/T 7714",
        "description": "中国国家标准《信息与文献 参考文献著录规则》",
        "citation_style": "gbt7714",
        "font_body": "SimSun",
        "font_heading": "SimHei",
        "font_size_body": 12,
        "font_size_heading": 16,
        "line_spacing": 1.5,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 3.17,
        "margin_right": 3.17,
    },
    "apa": {
        "id": "apa",
        "name": "APA 7th Edition",
        "description": "American Psychological Association 第7版格式",
        "citation_style": "apa",
        "font_body": "Times New Roman",
        "font_heading": "Times New Roman",
        "font_size_body": 12,
        "font_size_heading": 14,
        "line_spacing": 2.0,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 2.54,
        "margin_right": 2.54,
    },
    "ieee": {
        "id": "ieee",
        "name": "IEEE",
        "description": "Institute of Electrical and Electronics Engineers 格式",
        "citation_style": "ieee",
        "font_body": "Times New Roman",
        "font_heading": "Times New Roman",
        "font_size_body": 10,
        "font_size_heading": 14,
        "line_spacing": 1.0,
        "margin_top": 1.91,
        "margin_bottom": 2.54,
        "margin_left": 1.78,
        "margin_right": 1.78,
    },
}


class TemplateService:
    def __init__(self):
        self._docx_available = False
        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed, template generation unavailable")

    @property
    def available(self) -> bool:
        return self._docx_available

    def list_templates(self) -> List[Dict[str, Any]]:
        return [
            {
                "id": t["id"],
                "name": t["name"],
                "description": t["description"],
                "citation_style": t["citation_style"],
            }
            for t in BUILTIN_TEMPLATES.values()
        ]

    def generate_template(self, config: Dict[str, Any], output_path: str) -> str:
        if not self._docx_available:
            raise RuntimeError("python-docx 未安装，无法生成模板")

        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        template_id = config.get("template_id", "gbt7714")
        base = BUILTIN_TEMPLATES.get(template_id, BUILTIN_TEMPLATES["gbt7714"])

        font_body = config.get("font_body", base["font_body"])
        font_heading = config.get("font_heading", base["font_heading"])
        font_size_body = config.get("font_size_body", base["font_size_body"])
        font_size_heading = config.get("font_size_heading", base["font_size_heading"])
        line_spacing = config.get("line_spacing", base["line_spacing"])
        margin_top = config.get("margin_top", base["margin_top"])
        margin_bottom = config.get("margin_bottom", base["margin_bottom"])
        margin_left = config.get("margin_left", base["margin_left"])
        margin_right = config.get("margin_right", base["margin_right"])

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = font_body
        font.size = Pt(font_size_body)
        pf = style.paragraph_format
        pf.line_spacing = line_spacing
        pf.space_after = Pt(6)

        for level, size_ratio in [(1, 1.6), (2, 1.3), (3, 1.15)]:
            heading_style_name = f'Heading {level}'
            if heading_style_name in doc.styles:
                hs = doc.styles[heading_style_name]
                hs.font.name = font_heading
                hs.font.size = Pt(int(font_size_heading * size_ratio / 1.6))
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(6)

        sections = doc.sections
        if sections:
            section = sections[0]
            section.top_margin = Cm(margin_top)
            section.bottom_margin = Cm(margin_bottom)
            section.left_margin = Cm(margin_left)
            section.right_margin = Cm(margin_right)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("论文标题")
        title_run.font.name = font_heading
        title_run.font.size = Pt(font_size_heading + 4)
        title_run.bold = True

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run("作者姓名¹，合作者²")
        author_run.font.name = font_body
        author_run.font.size = Pt(font_size_body)

        doc.add_heading('摘要', level=1)
        doc.add_paragraph('在此输入摘要内容...')

        doc.add_heading('关键词', level=1)
        doc.add_paragraph('关键词1；关键词2；关键词3')

        doc.add_heading('1 引言', level=1)
        doc.add_paragraph('在此输入引言内容...')

        doc.add_heading('2 方法', level=1)
        doc.add_paragraph('在此输入方法内容...')

        doc.add_heading('3 结果', level=1)
        doc.add_paragraph('在此输入结果内容...')

        doc.add_heading('4 讨论', level=1)
        doc.add_paragraph('在此输入讨论内容...')

        doc.add_heading('参考文献', level=1)
        doc.add_paragraph('[1] 作者. 标题[J]. 期刊名, 年, 卷(期): 页码.')

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        logger.info("Template generated", path=str(output))
        return str(output)

    async def export_markdown_to_docx(
        self,
        markdown: str,
        template_id: str = "gbt7714",
        title: str = "document",
    ) -> bytes:
        if not self._docx_available:
            raise RuntimeError("python-docx 未安装，无法导出")

        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile
        import re

        base = BUILTIN_TEMPLATES.get(template_id, BUILTIN_TEMPLATES["gbt7714"])

        template_path = TEMPLATES_DIR / f"{template_id}.docx"
        if not template_path.exists():
            self.generate_template(
                {"template_id": template_id},
                str(template_path),
            )

        doc = Document(str(template_path))

        for line in markdown.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                level = min(level, 6)
                text = heading_match.group(2)
                doc.add_heading(text, level=level)
                continue

            if stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
                continue

            if re.match(r'^\d+\.\s+', stripped):
                content = re.sub(r'^\d+\.\s+', '', stripped)
                doc.add_paragraph(content, style='List Number')
                continue

            bold_text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            italic_text = re.sub(r'\*(.+?)\*', r'\1', bold_text)
            doc.add_paragraph(italic_text)

        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


_template_service: Optional[TemplateService] = None


def get_template_service() -> TemplateService:
    global _template_service
    if _template_service is None:
        _template_service = TemplateService()
    return _template_service
